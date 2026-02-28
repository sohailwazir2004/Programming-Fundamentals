import subprocess
import os
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime

BASE = os.path.dirname(os.path.abspath(__file__))

def add_code_block(doc, code, font_size=9):
    """Add formatted code block to document"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    # Set shading
    shading = p.paragraph_format.element
    return p

def add_output_block(doc, output, font_size=9):
    """Add output block"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(output)
    run.font.name = 'Consolas'
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p

def compile_and_run(java_file, lab_dir, input_data=None):
    """Compile and run a Java file, return output"""
    class_name = os.path.splitext(os.path.basename(java_file))[0]

    # Compile
    result = subprocess.run(
        ['javac', java_file],
        capture_output=True, text=True, cwd=lab_dir, timeout=30
    )
    if result.returncode != 0:
        return f"COMPILATION ERROR:\n{result.stderr}"

    # Run
    try:
        result = subprocess.run(
            ['java', class_name],
            capture_output=True, text=True, cwd=lab_dir,
            input=input_data, timeout=10
        )
        output = result.stdout
        if result.stderr:
            output += "\n" + result.stderr
        return output.strip()
    except subprocess.TimeoutExpired:
        return "Program timed out (requires interactive input)"

def create_lab_doc(lab_num, lab_title, activities, graded_tasks, output_path):
    """Create Word document for a lab"""
    doc = Document()

    # Title
    title = doc.add_heading(f'Lab {lab_num:02d}: {lab_title}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('')

    # Activities Section
    if activities:
        doc.add_heading('Solved Lab Activities', level=1)
        for act_name, code, output in activities:
            doc.add_heading(act_name, level=2)

            doc.add_heading('Source Code:', level=3)
            add_code_block(doc, code)

            doc.add_heading('Output:', level=3)
            add_output_block(doc, output)
            doc.add_paragraph('')  # spacing

    # Graded Tasks Section
    if graded_tasks:
        doc.add_heading('Graded Lab Tasks', level=1)
        for task_name, code, output in graded_tasks:
            doc.add_heading(task_name, level=2)

            doc.add_heading('Source Code:', level=3)
            add_code_block(doc, code)

            doc.add_heading('Output:', level=3)
            add_output_block(doc, output)
            doc.add_paragraph('')

    doc.save(output_path)
    print(f"Created: {output_path}")

def process_lab(lab_num, lab_title, lab_dir, activity_files, graded_files, activity_inputs, graded_inputs):
    """Process a complete lab"""
    activities = []
    for i, (fname, input_data) in enumerate(zip(activity_files, activity_inputs)):
        filepath = os.path.join(lab_dir, fname)
        if not os.path.exists(filepath):
            continue
        with open(filepath, 'r') as f:
            code = f.read()
        output = compile_and_run(filepath, lab_dir, input_data)
        activities.append((f"Activity {i+1}", code, output))
        print(f"  Lab {lab_num} Activity {i+1}: Done")

    graded = []
    for i, (fname, input_data) in enumerate(zip(graded_files, graded_inputs)):
        filepath = os.path.join(lab_dir, fname)
        if not os.path.exists(filepath):
            continue
        with open(filepath, 'r') as f:
            code = f.read()
        output = compile_and_run(filepath, lab_dir, input_data)
        graded.append((f"Graded Lab Task {i+1}", code, output))
        print(f"  Lab {lab_num} Graded Task {i+1}: Done")

    output_path = os.path.join(BASE, f"Lab{lab_num:02d}_{lab_title.replace(' ', '_')}.docx")
    create_lab_doc(lab_num, lab_title, activities, graded, output_path)

# ============ LAB 1 ============
print("Processing Lab 1...")
process_lab(
    1, "Java Installation Algorithms Errors Testing",
    os.path.join(BASE, "Lab1"),
    ["Payroll.java", "SalesTax.java"],
    [],
    ["40\n20\n", "100\n"],  # inputs for Payroll and SalesTax
    []
)

# ============ LAB 2 ============
print("Processing Lab 2...")
process_lab(
    2, "Java Fundamentals-I",
    os.path.join(BASE, "Lab2"),
    ["Activity1.java", "Activity2.java", "Activity3.java", "Activity4.java", "Activity5.java"],
    ["GLabTask1.java", "GLabTask2.java", "GLabTask3.java"],
    [
        None,           # Activity1 - no input
        "23 7\n",       # Activity2
        "Sheila Mann 23 120.5\n",  # Activity3
        None,           # Activity4
        "1\n",          # Activity5
    ],
    [
        None,                          # GLabTask1 - no input
        "10\n5\n",                     # GLabTask2
        "13 28\nMustafa\n48.30\n",     # GLabTask3
    ]
)

# ============ LAB 3 ============
print("Processing Lab 3...")
process_lab(
    3, "Java Fundamentals-II",
    os.path.join(BASE, "Lab3"),
    ["Activity1.java", "Activity2.java", "Activity3.java", "Activity4.java", "Activity5.java"],
    ["GLabTask1.java", "GLabTask2.java", "GLabTask3.java", "GLabTask4.java",
     "GLabTask5.java", "GLabTask6.java", "GLabTask7.java", "GLabTask8.java",
     "GLabTask9.java", "GLabTask10.java"],
    [
        None,           # Activity1
        None,           # Activity2
        None,           # Activity3
        None,           # Activity4
        "197.55\n",     # Activity5
    ],
    [
        "11.56\n",          # GLabTask1
        "6\n50\n",          # GLabTask2
        "28\n19\n18\n",     # GLabTask3
        "150\n",            # GLabTask4
        "100\n",            # GLabTask5
        "15.50\n40\n",      # GLabTask6
        "100\n200\n150\n50\n",  # GLabTask7
        "932\n",            # GLabTask8
        None,               # GLabTask9
        None,               # GLabTask10
    ]
)

print("\nAll done!")

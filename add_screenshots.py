import subprocess
import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS_DIR = os.path.join(BASE, "screenshots")
os.makedirs(SCREENSHOTS_DIR, exist_ok=True)

def compile_and_run(java_file, lab_dir, input_data=None):
    class_name = os.path.splitext(os.path.basename(java_file))[0]
    result = subprocess.run(['javac', java_file], capture_output=True, text=True, cwd=lab_dir, timeout=30)
    if result.returncode != 0:
        return f"COMPILATION ERROR:\n{result.stderr}"
    try:
        result = subprocess.run(['java', class_name], capture_output=True, text=True, cwd=lab_dir, input=input_data, timeout=10)
        output = result.stdout
        if result.stderr:
            output += "\n" + result.stderr
        return output.strip()
    except subprocess.TimeoutExpired:
        return "Program timed out"

def create_output_html(code, output, title):
    """Create HTML for screenshot"""
    escaped_code = code.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    escaped_output = output.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return f"""<!DOCTYPE html>
<html><head><style>
body {{ font-family: 'Segoe UI', sans-serif; margin: 20px; background: #f5f5f5; max-width: 800px; }}
.container {{ background: white; border-radius: 8px; padding: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.1); margin-bottom: 20px; }}
h2 {{ color: #1a1a2e; border-bottom: 2px solid #e94560; padding-bottom: 8px; font-size: 16px; }}
pre {{ background: #1e1e2e; color: #cdd6f4; padding: 15px; border-radius: 6px; font-size: 13px; overflow-x: auto; line-height: 1.5; }}
.output pre {{ background: #0d1117; color: #58a6ff; }}
.title {{ font-size: 14px; color: #666; margin-bottom: 5px; }}
</style></head><body>
<div class="container">
<h2>Source Code - {title}</h2>
<pre>{escaped_code}</pre>
</div>
<div class="container output">
<h2>Output</h2>
<pre>{escaped_output}</pre>
</div>
</body></html>"""

def create_lab_doc_with_screenshots(lab_num, lab_title, entries, output_path):
    doc = Document()
    title = doc.add_heading(f'Lab {lab_num:02d}: {lab_title}', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')

    has_activities = any(e[0].startswith("Activity") for e in entries)
    has_graded = any(e[0].startswith("Graded") for e in entries)

    if has_activities:
        doc.add_heading('Solved Lab Activities', level=1)

    switched = False
    for entry_name, code, output, screenshot_path in entries:
        if not switched and entry_name.startswith("Graded"):
            switched = True
            doc.add_heading('Graded Lab Tasks', level=1)

        doc.add_heading(entry_name, level=2)

        doc.add_heading('Source Code:', level=3)
        p = doc.add_paragraph()
        run = p.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

        doc.add_heading('Output:', level=3)
        p = doc.add_paragraph()
        run = p.add_run(output)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 100, 0)

        if screenshot_path and os.path.exists(screenshot_path):
            doc.add_heading('Screenshot:', level=3)
            doc.add_picture(screenshot_path, width=Inches(6))

        doc.add_paragraph('')

    doc.save(output_path)
    print(f"  Saved: {output_path}")

# Collect all data
labs = [
    {
        "num": 1, "title": "Java Installation Algorithms Errors Testing",
        "dir": "Lab1",
        "files": [
            ("Activity 1 - Payroll Program", "Payroll.java", "40\n20\n"),
            ("Activity 2 - SalesTax Debugging", "SalesTax.java", "100\n"),
        ]
    },
    {
        "num": 2, "title": "Java Fundamentals-I",
        "dir": "Lab2",
        "files": [
            ("Activity 1 - Variable Manipulation", "Activity1.java", None),
            ("Activity 2 - Input Statements", "Activity2.java", "23 7\n"),
            ("Activity 3 - Reading Strings & Numbers", "Activity3.java", "Sheila Mann 23 120.5\n"),
            ("Activity 4 - Characters & Integers", "Activity4.java", None),
            ("Activity 5 - Constants (Circle Area)", "Activity5.java", "1\n"),
            ("Graded Lab Task 1 - Average Calculation", "GLabTask1.java", None),
            ("Graded Lab Task 2 - Rectangle Area & Perimeter", "GLabTask2.java", "10\n5\n"),
            ("Graded Lab Task 3 - Comprehensive Program", "GLabTask3.java", "13 28\nMustafa\n48.30\n"),
        ]
    },
    {
        "num": 3, "title": "Java Fundamentals-II",
        "dir": "Lab3",
        "files": [
            ("Activity 1 - E Notation", "Activity1.java", None),
            ("Activity 2 - Hourly Wages + Overtime", "Activity2.java", None),
            ("Activity 3 - Retirement Contributions", "Activity3.java", None),
            ("Activity 4 - Formatted Output (printf)", "Activity4.java", None),
            ("Activity 5 - Sales Tax Calculation", "Activity5.java", "197.55\n"),
            ("Graded Lab Task 1 - Money Change", "GLabTask1.java", "11.56\n"),
            ("Graded Lab Task 2 - Apple Distribution", "GLabTask2.java", "6\n50\n"),
            ("Graded Lab Task 3 - Desks Calculation", "GLabTask3.java", "28\n19\n18\n"),
            ("Graded Lab Task 4 - Digital Clock", "GLabTask4.java", "150\n"),
            ("Graded Lab Task 5 - Milk Cartons", "GLabTask5.java", "100\n"),
            ("Graded Lab Task 6 - Summer Job Income", "GLabTask6.java", "15.50\n40\n"),
            ("Graded Lab Task 7 - Cricket Tickets", "GLabTask7.java", "100\n200\n150\n50\n"),
            ("Graded Lab Task 8 - Sum of Digits", "GLabTask8.java", "932\n"),
            ("Graded Lab Task 9 - Printf Output", "GLabTask9.java", None),
            ("Graded Lab Task 10 - Trig Table", "GLabTask10.java", None),
        ]
    }
]

# Generate HTML files for screenshots
all_html_files = []
for lab in labs:
    lab_dir = os.path.join(BASE, lab["dir"])
    for entry_name, fname, input_data in lab["files"]:
        filepath = os.path.join(lab_dir, fname)
        with open(filepath, 'r') as f:
            code = f.read()
        output = compile_and_run(filepath, lab_dir, input_data)

        safe_name = f"Lab{lab['num']}_{fname.replace('.java','')}"
        html_path = os.path.join(SCREENSHOTS_DIR, f"{safe_name}.html")
        png_path = os.path.join(SCREENSHOTS_DIR, f"{safe_name}.png")

        html_content = create_output_html(code, output, entry_name)
        with open(html_path, 'w') as f:
            f.write(html_content)

        all_html_files.append({
            "lab_num": lab["num"],
            "entry_name": entry_name,
            "code": code,
            "output": output,
            "html_path": html_path,
            "png_path": png_path
        })

# Save manifest for screenshot taking
manifest_path = os.path.join(SCREENSHOTS_DIR, "manifest.json")
with open(manifest_path, 'w') as f:
    json.dump(all_html_files, f, indent=2)

print(f"Generated {len(all_html_files)} HTML files for screenshots.")
print(f"Manifest saved to: {manifest_path}")
print("Now take screenshots with Playwright, then run generate_final_docs.py")
